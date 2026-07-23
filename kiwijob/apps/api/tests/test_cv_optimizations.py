from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from conftest import auth_headers


def _docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Test Candidate", 0)
    document.add_paragraph("Data analyst with Python and SQL experience.")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_create_edit_and_download_cv_optimization() -> None:
    with TestClient(app) as client:
        headers, _ = auth_headers(client)
        resume_response = client.post(
            "/resumes/upload",
            headers=headers,
            files={
                "file": (
                    "candidate.docx",
                    _docx_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert resume_response.status_code == 200
        application_response = client.post(
            "/jobs/save",
            headers=headers,
            json={
                "title": "Data Analyst",
                "company": "Example",
                "description": "Python, SQL, reporting and stakeholder communication.",
                "url": "https://example.test/jobs/data-analyst",
                "source_website": "example.test",
            },
        )
        assert application_response.status_code == 200

        created = client.post(
            "/cv-optimizations",
            headers=headers,
            json={
                "application_id": application_response.json()["id"],
                "resume_id": resume_response.json()["id"],
            },
        )
        assert created.status_code == 201
        body = created.json()
        assert body["optimized_text"]
        assert body["suggestions"]
        assert body["resume_id"] == resume_response.json()["id"]

        edited = client.put(
            f"/cv-optimizations/{body['id']}",
            headers=headers,
            json={"title": "Data Analyst CV", "optimized_text": body["optimized_text"] + "\nVerified edit"},
        )
        assert edited.status_code == 200
        assert edited.json()["title"] == "Data Analyst CV"

        download = client.get(f"/cv-optimizations/{body['id']}/download", headers=headers)
        assert download.status_code == 200
        assert download.headers["content-type"].startswith(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        assert download.content[:2] == b"PK"

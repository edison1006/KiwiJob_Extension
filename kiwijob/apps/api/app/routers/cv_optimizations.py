from __future__ import annotations

from datetime import datetime
from io import BytesIO
import re

from docx import Document
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import selectinload
from sqlmodel import Session, select

from app.db.session import get_session
from app.deps import get_current_user
from app.models import Application, CvOptimization, Resume, User
from app.schemas import CvOptimizationCreateIn, CvOptimizationOut, CvOptimizationUpdateIn
from app.services.cv_optimize import optimize_cv

router = APIRouter(prefix="/cv-optimizations", tags=["cv-optimizations"])


def _out(row: CvOptimization) -> CvOptimizationOut:
    return CvOptimizationOut.model_validate(row, from_attributes=True)


def _owned(session: Session, optimization_id: int, user_id: int) -> CvOptimization:
    row = session.exec(
        select(CvOptimization).where(CvOptimization.id == optimization_id, CvOptimization.user_id == user_id)
    ).first()
    if not row:
        raise HTTPException(status_code=404, detail="CV optimization not found")
    return row


@router.post("", response_model=CvOptimizationOut, status_code=201)
def create_optimization(
    body: CvOptimizationCreateIn,
    session: Session = Depends(get_session),
    user: User = Depends(get_current_user),
):
    assert user.id is not None
    resume = session.exec(
        select(Resume).where(Resume.id == body.resume_id, Resume.user_id == user.id)
    ).first()
    application = session.exec(
        select(Application)
        .where(Application.id == body.application_id, Application.user_id == user.id)
        .options(selectinload(Application.job_post))
    ).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    if not application or not application.job_post:
        raise HTTPException(status_code=404, detail="Application not found")
    if not resume.extracted_text.strip():
        raise HTTPException(status_code=400, detail="The selected CV has no extractable text")
    jd = "\n\n".join(
        part for part in (application.job_post.title, application.job_post.description or "", application.job_post.visa_requirement or "") if part
    )
    score, suggestions, optimized_text = optimize_cv(resume.extracted_text, jd)
    row = CvOptimization(
        user_id=user.id,
        resume_id=resume.id,
        application_id=application.id,
        title=f"{application.job_post.title} — Optimized CV",
        match_score=score,
        suggestions=[item.model_dump() for item in suggestions],
        optimized_text=optimized_text,
    )
    session.add(row)
    session.commit()
    session.refresh(row)
    return _out(row)


@router.get("", response_model=list[CvOptimizationOut])
def list_optimizations(session: Session = Depends(get_session), user: User = Depends(get_current_user)):
    assert user.id is not None
    rows = session.exec(
        select(CvOptimization).where(CvOptimization.user_id == user.id).order_by(CvOptimization.updated_at.desc())
    ).all()
    return [_out(row) for row in rows]


@router.put("/{optimization_id}", response_model=CvOptimizationOut)
def update_optimization(
    optimization_id: int,
    body: CvOptimizationUpdateIn,
    session: Session = Depends(get_session),
    user: User = Depends(get_current_user),
):
    assert user.id is not None
    row = _owned(session, optimization_id, user.id)
    if body.title is not None:
        row.title = body.title.strip()
    if body.optimized_text is not None:
        row.optimized_text = body.optimized_text
    if body.suggestions is not None:
        row.suggestions = [item.model_dump() for item in body.suggestions]
    row.updated_at = datetime.utcnow()
    session.add(row)
    session.commit()
    session.refresh(row)
    return _out(row)


@router.get("/{optimization_id}/download")
def download_optimization(
    optimization_id: int,
    session: Session = Depends(get_session),
    user: User = Depends(get_current_user),
):
    assert user.id is not None
    row = _owned(session, optimization_id, user.id)
    document = Document()
    document.add_heading(row.title, 0)
    for raw_line in row.optimized_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if len(line) < 70 and (line.isupper() or line.endswith(":")):
            document.add_heading(line.rstrip(":"), level=1)
        elif line.startswith(("- ", "• ")):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(line)
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    filename = re.sub(r"[^A-Za-z0-9_-]+", "-", row.title).strip("-") or "optimized-cv"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}.docx"'},
    )
